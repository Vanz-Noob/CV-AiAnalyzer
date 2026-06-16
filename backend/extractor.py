"""
Ekstraktor teks CV — Mendukung format PDF dan DOC/DOCX.

Menggunakan PyMuPDF (fitz) untuk PDF dan python-docx untuk DOCX.
Menangani ekstraksi teks dari paragraf dan tabel (untuk DOCX).
"""

import fitz
from docx import Document


def extract_text_from_pdf(content: bytes) -> str:
    """Ekstrak seluruh teks dari file PDF.

    Args:
        content: Konten biner file PDF.

    Returns:
        String teks gabungan dari semua halaman, dipisahkan newline.
    """
    doc = fitz.open(stream=content, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts).strip()


def extract_text_from_docx(content: bytes) -> str:
    """Ekstrak teks dari file DOCX, termasuk isi tabel.

    Args:
        content: Konten biner file DOCX.

    Returns:
        String teks gabungan dari paragraf dan baris tabel,
        sel-sel tabel dipisahkan dengan " | ".
    """
    from io import BytesIO

    doc = Document(BytesIO(content))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_parts.append(row_text)
    return "\n".join(text_parts).strip()


def extract_cv_text(content: bytes, filename: str) -> str:
    """Ekstrak teks dari file CV berdasarkan ekstensi filename.

    Args:
        content: Konten biner file CV.
        filename: Nama file (digunakan untuk menentukan format).

    Returns:
        String teks yang berhasil diekstrak.

    Raises:
        ValueError: Jika format file tidak didukung (bukan PDF/DOC/DOCX).
    """
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return extract_text_from_pdf(content)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file format: .{ext}")
